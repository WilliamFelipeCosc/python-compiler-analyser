import pandas as pd
from fpdf import FPDF
from docx import Document
from lexical_analysis import analisar_lexico

def gerar_pdf_tabela(tokens, nome_arquivo_pdf="tabela_tokens.pdf"):
    # Cria DataFrame
    df = pd.DataFrame(tokens, columns=["Nome", "Token", "Tipo", "Descrição", "Linha"])

    # Cria PDF
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=10)
    col_widths = [30, 60, 30, 60, 15]

    # Cabeçalho
    headers = ["Nome", "Token", "Tipo", "Descrição", "Linha"]
    for i, header in enumerate(headers):
        pdf.cell(col_widths[i], 10, header, border=1)
    pdf.ln()

    # Linhas da tabela
    for _, row in df.iterrows():
        pdf.cell(col_widths[0], 10, str(row["Nome"]), border=1)
        pdf.cell(col_widths[1], 10, str(row["Token"])[:30], border=1)  # Limita tamanho
        pdf.cell(col_widths[2], 10, str(row["Tipo"]), border=1)
        pdf.cell(col_widths[3], 10, str(row["Descrição"])[:30], border=1)
        pdf.cell(col_widths[4], 10, str(row["Linha"]), border=1)
        pdf.ln()

    pdf.output(nome_arquivo_pdf)
    print(f"Tabela salva em {nome_arquivo_pdf}")

def gerar_docx_tabela(tokens, nome_arquivo_docx="tabela_tokens.docx"):
    # Cria DataFrame
    df = pd.DataFrame(tokens, columns=["Nome", "Token", "Tipo", "Descrição", "Linha"])

    # Cria documento Word
    doc = Document()
    doc.add_heading('Tabela de Tokens', 0)
    table = doc.add_table(rows=1, cols=len(df.columns))
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = col

    # Adiciona linhas à tabela
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, item in enumerate(row):
            row_cells[i].text = str(item)

    doc.save(nome_arquivo_docx)
    print(f"Tabela salva em {nome_arquivo_docx}")

if __name__ == "__main__":
    with open("main.py", "r", encoding="utf-8") as f:
        codigo = f.read()
    tokens = analisar_lexico(codigo)
    gerar_pdf_tabela(tokens)
    gerar_docx_tabela(tokens)